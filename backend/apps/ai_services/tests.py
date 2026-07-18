import json
from pathlib import Path
import re
from types import SimpleNamespace
from tempfile import TemporaryDirectory
from unittest.mock import patch

import fitz
from django.test import SimpleTestCase, override_settings
from docx import Document

from .resume_preprocessor import (
    cleanup_punctuation,
    coerce_text,
    normalize_tokens,
    normalize_whitespace,
    preprocess_for_matching,
    preprocess_for_semantic_matching,
    safe_lower,
)
from .resume_text_extractor import ResumeTextExtractionError, extract_resume_text
from .resume_screening import (
    build_resume_screening,
    calculate_education_score,
    calculate_experience_score,
    calculate_skill_score,
    extract_education,
    extract_experience,
)
from .scoring import calculate_final_score, calculate_score_breakdown
from .exceptions import AIServiceUnavailable
from .semantic_matcher import semantic_similarity
from .skill_extractor import (
    _load_spacy_model,
    extract_skill_labels,
    extract_skills,
    get_skill_display_labels,
    normalize_skill_key,
    normalize_text,
)


class ResumeTextExtractorTests(SimpleTestCase):
    def test_extract_resume_text_from_pdf(self):
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'resume.pdf'
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), 'Python developer')
            document.save(file_path)
            document.close()

            self.assertEqual(extract_resume_text(file_path), 'Python developer')

    def test_extract_resume_text_from_docx_includes_paragraphs_and_table_cells(self):
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'resume.docx'
            document = Document()
            document.add_paragraph('Backend engineer')
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = 'Django'
            document.save(file_path)

            self.assertEqual(extract_resume_text(file_path), 'Backend engineer\nDjango')

    def test_extract_resume_text_rejects_unsupported_file_type(self):
        with self.assertRaisesMessage(ResumeTextExtractionError, 'Unsupported resume file type: .txt'):
            extract_resume_text('resume.txt')

    def test_extract_resume_text_rejects_missing_local_file(self):
        with self.assertRaisesMessage(ResumeTextExtractionError, 'Resume file does not exist'):
            extract_resume_text('missing-resume.pdf')


class ResumePreprocessorTests(SimpleTestCase):
    def test_coerce_text_handles_none_strings_and_non_string_input(self):
        self.assertEqual(coerce_text(None), '')
        self.assertEqual(coerce_text('Resume'), 'Resume')
        self.assertEqual(coerce_text(2026), '2026')

    def test_normalize_whitespace_collapses_repeated_whitespace(self):
        self.assertEqual(normalize_whitespace('  Python\n\t  Django   developer  '), 'Python Django developer')

    def test_safe_lower_lowercases_without_other_cleanup(self):
        self.assertEqual(safe_lower('  React.JS!  '), '  react.js!  ')

    def test_cleanup_punctuation_preserves_common_skill_symbols_by_default(self):
        self.assertEqual(cleanup_punctuation('C++, C#, React.js, Node.js!'), 'C++ C# React.js Node.js ')

    def test_normalize_tokens_can_remove_skill_symbols_when_requested(self):
        self.assertEqual(
            normalize_tokens('C++, C#, React.js', preserve_skill_symbols=False),
            'c c react js',
        )

    def test_preprocess_for_matching_returns_safe_normalized_matching_text(self):
        self.assertEqual(preprocess_for_matching('  Python,   React.js!  '), 'python react.js')

    def test_preprocess_for_semantic_matching_keeps_normalized_copy_only(self):
        original_text = '  Backend   Engineer: Python/Django  '

        self.assertEqual(preprocess_for_semantic_matching(original_text), 'backend engineer python django')
        self.assertEqual(original_text, '  Backend   Engineer: Python/Django  ')


class _FakeStrings:
    def __getitem__(self, match_id):
        return match_id


class _FakeVocab:
    strings = _FakeStrings()


class _FakeNLP:
    vocab = _FakeVocab()

    def __call__(self, text):
        return text

    def make_doc(self, text):
        return text


class _FakePhraseMatcher:
    def __init__(self, _vocab, attr=None):
        self.patterns_by_skill = {}

    def add(self, skill_key, patterns):
        self.patterns_by_skill[skill_key] = patterns

    def __call__(self, doc):
        matches = []
        for skill_key, patterns in self.patterns_by_skill.items():
            if any(_fake_contains_alias(doc, pattern) for pattern in patterns):
                matches.append((skill_key, 0, 1))
        return matches


def _fake_contains_alias(normalized_text, normalized_alias):
    pattern = rf'(?<![a-z0-9]){re.escape(normalized_alias)}(?![a-z0-9])'
    return bool(re.search(pattern, normalized_text))


class SkillExtractorTests(SimpleTestCase):
    def test_normalize_text_lowercases_and_removes_extra_punctuation(self):
        self.assertEqual(normalize_text('  Python,   React.js!  '), 'python react.js')

    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_extract_skills_normalizes_aliases_and_returns_canonical_names(self, _mock_spacy_model, _mock_phrase_matcher_class):
        resume_text = 'Built RESTful APIs with Python, Django, ReactJS, PostgreSQL, AWS, and C++.'

        self.assertEqual(
            extract_skills(resume_text),
            ['aws', 'c++', 'django', 'postgresql', 'python', 'react', 'rest api'],
        )

    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_extract_skills_does_not_match_alias_inside_another_word(self, _mock_spacy_model, _mock_phrase_matcher_class):
        self.assertEqual(extract_skills('Enjoys javascript.', {'java': ('java',)}), [])

    def test_extract_skills_accepts_an_empty_custom_dictionary(self):
        self.assertEqual(extract_skills('Python', {}), [])

    @patch('apps.ai_services.skill_extractor._load_spacy_model', side_effect=AIServiceUnavailable('spaCy unavailable'))
    def test_extract_skills_uses_dictionary_fallback_when_spacy_unavailable(self, _mock_spacy_model):
        self.assertEqual(
            extract_skills('Built RESTful APIs with py, js, nodejs, postgres, and reactjs.'),
            ['javascript', 'node.js', 'postgresql', 'python', 'react', 'rest api'],
        )

    def test_load_spacy_model_raises_when_spacy_dependency_import_fails(self):
        _load_spacy_model.cache_clear()
        self.addCleanup(_load_spacy_model.cache_clear)

        with (
            patch('apps.ai_services.skill_extractor.importlib.util.find_spec', return_value=object()),
            patch(
                'apps.ai_services.skill_extractor.importlib.import_module',
                side_effect=ModuleNotFoundError('click'),
            ) as mock_import_module,
        ):
            with self.assertRaises(AIServiceUnavailable):
                _load_spacy_model()

        mock_import_module.assert_called_once_with('spacy')

    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_extract_skill_labels_uses_spacy_matches(self, _mock_spacy_model, _mock_phrase_matcher_class):
        with patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher), patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP()):
            self.assertEqual(extract_skill_labels('py js reactjs nodejs postgres'), [
                'JavaScript',
                'Node.js',
                'PostgreSQL',
                'Python',
                'React',
            ])

    def test_normalize_skill_key_maps_aliases_to_internal_keys(self):
        self.assertEqual(normalize_skill_key('py'), 'python')
        self.assertEqual(normalize_skill_key('js'), 'javascript')
        self.assertEqual(normalize_skill_key('reactjs'), 'react')
        self.assertEqual(normalize_skill_key('nodejs'), 'node.js')
        self.assertEqual(normalize_skill_key('postgres'), 'postgresql')

    def test_skill_display_labels_are_additive_and_canonical(self):
        self.assertEqual(
            get_skill_display_labels(['python', 'javascript', 'react', 'node.js', 'postgresql']),
            ['Python', 'JavaScript', 'React', 'Node.js', 'PostgreSQL'],
        )
        with patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher), patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP()):
            self.assertEqual(extract_skill_labels('py js reactjs nodejs postgres'), [
                'JavaScript',
                'Node.js',
                'PostgreSQL',
                'Python',
                'React',
            ])

    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_extract_skills_uses_spacy_phrase_matcher_when_available(
        self, _mock_spacy_model, _mock_phrase_matcher_class
    ):
        self.assertEqual(
            extract_skills('Experience with PY, ReactJS, NodeJS, and Postgres.'),
            ['node.js', 'postgresql', 'python', 'react'],
        )


class SemanticMatcherTests(SimpleTestCase):
    def test_semantic_similarity_uses_fast_fallback_by_default(self):
        self.assertEqual(semantic_similarity('Python Django developer', 'Django Python engineer'), 66.67)

    @override_settings(AI_USE_SENTENCE_BERT=True)
    @patch('apps.ai_services.semantic_matcher._get_model', side_effect=ModuleNotFoundError)
    def test_semantic_similarity_falls_back_when_dependency_is_unavailable(self, _mock_model):
        self.assertEqual(semantic_similarity('Python Django developer', 'Django Python engineer'), 66.67)

    @override_settings(AI_USE_SENTENCE_BERT=True)
    @patch('apps.ai_services.semantic_matcher._get_model', side_effect=OSError('offline model download failed'))
    def test_semantic_similarity_falls_back_when_model_loading_fails(self, _mock_model):
        self.assertEqual(semantic_similarity('Python Django developer', 'Django Python engineer'), 66.67)

    @override_settings(AI_USE_SENTENCE_BERT=True)
    @patch('apps.ai_services.semantic_matcher._get_model')
    def test_semantic_similarity_falls_back_when_encoding_fails(self, mock_get_model):
        mock_get_model.return_value.encode.side_effect = RuntimeError('tensor execution failed')

        self.assertEqual(semantic_similarity('Python Django developer', 'Django Python engineer'), 66.67)

    @override_settings(AI_USE_SENTENCE_BERT=True)
    @patch('apps.ai_services.semantic_matcher._get_model')
    def test_semantic_similarity_falls_back_when_tensor_handling_fails(self, mock_get_model):
        mock_get_model.return_value.encode.return_value = []

        self.assertEqual(semantic_similarity('Python Django developer', 'Django Python engineer'), 66.67)

    @override_settings(AI_USE_SENTENCE_BERT=True)
    @patch('apps.ai_services.semantic_matcher._get_model')
    def test_semantic_similarity_uses_model_embeddings_when_enabled_and_available(self, mock_get_model):
        mock_get_model.return_value.encode.return_value = [_Vector(0.75), _Vector(0.75)]

        self.assertEqual(semantic_similarity('  Python,   developer!  ', 'Backend\nengineer'), 75.0)
        mock_get_model.return_value.encode.assert_called_once_with(
            ['python developer', 'backend engineer'],
            convert_to_tensor=True,
            normalize_embeddings=True,
        )

    @override_settings(AI_USE_SENTENCE_BERT=True)
    @patch('apps.ai_services.semantic_matcher._get_model')
    def test_semantic_similarity_normalizes_model_scores_to_zero_to_one_hundred(self, mock_get_model):
        mock_get_model.return_value.encode.return_value = [_Vector(1.5), _Vector(1.5)]
        self.assertEqual(semantic_similarity('Python', 'Python'), 100.0)

        mock_get_model.return_value.encode.return_value = [_Vector(-0.25), _Vector(-0.25)]
        self.assertEqual(semantic_similarity('Python', 'Python'), 0.0)

    def test_semantic_similarity_returns_zero_for_blank_input(self):
        self.assertEqual(semantic_similarity('', 'Backend engineer'), 0.0)
        self.assertEqual(semantic_similarity('Python developer', '   '), 0.0)


class _Vector:
    def __init__(self, similarity):
        self.similarity = similarity

    def __matmul__(self, _other):
        return self

    def item(self):
        return self.similarity


class ScoringTests(SimpleTestCase):
    def test_calculate_final_score_uses_required_weights(self):
        self.assertEqual(calculate_final_score(80, 70, 60, 50), 70.0)

    def test_calculate_final_score_matches_exact_required_formula(self):
        semantic_score = 82.5
        skill_score = 75.0
        experience_score = 60.0
        education_score = 80.0

        expected = round(
            (0.4 * semantic_score)
            + (0.3 * skill_score)
            + (0.2 * experience_score)
            + (0.1 * education_score),
            2,
        )

        self.assertEqual(
            calculate_final_score(semantic_score, skill_score, experience_score, education_score),
            expected,
        )

    def test_calculate_score_breakdown_returns_components_and_final_score(self):
        self.assertEqual(
            calculate_score_breakdown(80, 70, 60, 50),
            {
                'semantic_score': 80,
                'skill_score': 70,
                'experience_score': 60,
                'education_score': 50,
                'final_score': 70.0,
            },
        )

    def test_calculate_final_score_rejects_out_of_range_component(self):
        with self.assertRaisesMessage(ValueError, 'skill_score must be between 0 and 100'):
            calculate_final_score(80, 101, 60, 50)


class ResumeScreeningScoreComponentTests(SimpleTestCase):
    @patch('apps.ai_services.resume_screening.build_ml_screening_result')
    @patch('apps.ai_services.resume_screening.semantic_similarity', return_value=80.0)
    @patch('apps.ai_services.resume_screening.extract_resume_text')
    def test_build_resume_screening_uses_trained_ml_score_as_final_score(
        self, extract_resume_text, _semantic_similarity, build_ml_screening_result
    ):
        from apps.jobs.models import JobRequirement

        extract_resume_text.return_value = (
            "Skills: Python Django. Education: Bachelor's degree. "
            'Experience: Backend developer with 5 years of experience.'
        )
        build_ml_screening_result.return_value = {
            'ml_suitability_score': 88.25,
            'ml_match_label': 'strong_match',
            'ml_confidence': 0.92,
            'model_version': 'test-trained-model-v1',
        }
        requirements = [
            SimpleNamespace(
                requirement_type=JobRequirement.RequirementType.SKILL,
                description='Python and Django',
                weight_score=30,
            ),
            SimpleNamespace(
                requirement_type=JobRequirement.RequirementType.EXPERIENCE,
                description='At least 3 years of professional experience',
                weight_score=20,
            ),
            SimpleNamespace(
                requirement_type=JobRequirement.RequirementType.EDUCATION,
                description="Bachelor's degree",
                weight_score=10,
            ),
        ]
        application = SimpleNamespace(
            resume_id=None,
            applicant=SimpleNamespace(
                applicant_profile=SimpleNamespace(resume_file=SimpleNamespace(path='/tmp/resume.pdf'))
            ),
            job=SimpleNamespace(
                title='Backend Engineer',
                description='Build APIs',
                requirements=SimpleNamespace(all=lambda: requirements),
            ),
        )

        result = build_resume_screening(application)

        self.assertEqual(result['final_score'], 88.25)
        self.assertEqual(result['score_explanation']['final_score'], 88.25)
        self.assertEqual(result['score_explanation']['score_source'], 'trained_ml_model')
        self.assertEqual(result['score_explanation']['model_version'], 'test-trained-model-v1')
        self.assertEqual(result['score_explanation']['rule_based_score'], 92.0)
        self.assertEqual(
            result['score_explanation']['formula'],
            'final_score = trained_resume_match_model(feature_vector)',
        )
        build_ml_screening_result.assert_called_once()

    def test_extract_experience_uses_highest_explicit_year_value(self):
        result = extract_experience('2 years support and 5+ yrs development')

        self.assertEqual(result['years'], 5.0)
        self.assertIn('years', result)

    def test_extract_experience_preprocesses_text_before_matching(self):
        result = extract_experience('Worked with Python.\n  3+     YRS!!!')

        self.assertEqual(result['years'], 3.0)
        self.assertIn('3+ YRS', result['raw_mentions'])

    def test_extract_experience_detects_roles_companies_and_internships(self):
        result = extract_experience(
            'Software engineer at ABC Company. Worked as developer. Internship at Beta Labs.'
        )

        self.assertEqual(result['years'], 0.0)
        self.assertIn('software engineer', result['roles'])
        self.assertIn('developer', result['roles'])
        self.assertIn('intern', result['roles'])
        self.assertIn('ABC Company', result['companies'])
        self.assertIn('Beta Labs', result['companies'])
        self.assertIn('Internship at Beta Labs', result['internships'])
        self.assertIn('Software engineer at ABC Company', result['matched_phrases'])

    def test_extract_experience_keeps_old_expected_keys_while_returning_richer_object(self):
        result = extract_experience('5 yrs as a backend developer')

        self.assertIn('years', result)
        self.assertIn('roles', result)
        self.assertIn('companies', result)
        self.assertIn('internships', result)
        self.assertIn('matched_phrases', result)
        self.assertIn('raw_mentions', result)

    def test_extract_education_uses_highest_mentioned_level(self):
        result = extract_education("Bachelor's degree and master's degree")

        self.assertEqual(result['level'], 'master')
        self.assertEqual(result['level_label'], 'Master')

    def test_extract_education_preprocesses_text_before_matching(self):
        result = extract_education('Completed B.Sc, then MBA.')

        self.assertEqual(result['level'], 'master')
        self.assertIn('Bachelor', result['matched_keywords'])
        self.assertIn('Master', result['matched_keywords'])

    def test_extract_education_detects_all_supported_levels(self):
        examples = {
            'secondary': 'High School certificate',
            'diploma': 'Diploma in Information Technology',
            'associate': 'Associate Degree in Computer Science',
            'bachelor': 'Bachelor Degree in Software Engineering',
            'master': 'Master of Computer Science',
            'doctorate': 'PhD in Software Engineering',
        }

        for expected_level, text in examples.items():
            with self.subTest(expected_level=expected_level):
                self.assertEqual(extract_education(text)['level'], expected_level)

    def test_extract_education_detects_fields_of_study(self):
        result = extract_education(
            'Bachelor Degree in Computer Science and Diploma in Information Technology. '
            'Completed Software Engineering capstone.'
        )

        self.assertEqual(result['level'], 'bachelor')
        self.assertEqual(
            result['fields_of_study'],
            ['Computer Science', 'Software Engineering', 'Information Technology'],
        )
        self.assertIn('Degree', result['matched_keywords'])
        self.assertIn('Computer Science', result['raw_mentions'])

    def test_extract_education_keeps_old_expected_keys_while_returning_richer_object(self):
        result = extract_education('Bachelor Degree in Computer Science')

        self.assertIn('level', result)
        self.assertIn('level_label', result)
        self.assertIn('fields_of_study', result)
        self.assertIn('matched_keywords', result)
        self.assertIn('raw_mentions', result)

    def test_skill_score_calculates_required_skill_coverage(self):
        self.assertEqual(calculate_skill_score(['django', 'python'], ['django', 'python', 'sql']), 66.67)

    def test_skill_score_uses_requirement_weight_scores_when_available(self):
        skill_requirements = [
            {'skills': ['python'], 'weight_score': 80.0},
            {'skills': ['react'], 'weight_score': 20.0},
        ]

        self.assertEqual(calculate_skill_score(['python'], ['python', 'react'], skill_requirements), 80.0)

    def test_experience_score_is_capped_at_one_hundred(self):
        self.assertEqual(calculate_experience_score({'years': 5.0}, {'years': 3.0}), 100.0)

    def test_education_score_is_zero_when_required_level_is_missing_from_resume(self):
        self.assertEqual(calculate_education_score({'level': None}, {'level': 'bachelor'}), 0.0)


class LinkedInProfileImporterTests(SimpleTestCase):
    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_linkedin_profile_fixture_is_parsed_into_expected_sections(self, _mock_spacy_model, _mock_phrase_matcher_class):
        from .linkedin_profile_importer import build_linkedin_profile_import

        fixture_dir = Path(__file__).resolve().parents[1] / 'users' / 'test_fixtures'
        raw_text = (fixture_dir / 'linkedin_profile_sample_raw.txt').read_text()
        expected = json.loads((fixture_dir / 'linkedin_profile_sample_expected.json').read_text())

        parsed = build_linkedin_profile_import(raw_text)

        self.assertEqual(parsed['full_name'], expected['full_name'])
        self.assertEqual(parsed['headline'], expected['headline'])
        self.assertEqual(parsed['location'], expected['location'])
        self.assertEqual(parsed['linkedin_url'], expected['linkedin_url'])
        self.assertEqual(parsed['summary'], expected['summary'])
        self.assertEqual(parsed['skills'][:len(expected['skills'])], expected['skills'])
        for expected_skill in ['Java', 'Kubernetes', 'AWS']:
            self.assertIn(expected_skill, parsed['skills'])
        self.assertEqual(parsed['certifications'], expected['certifications'])
        self.assertEqual(parsed['experience'], expected['experience'])
        self.assertEqual(parsed['education'], expected['education'])


    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_linkedin_profile_parser_merges_sidebar_and_headline_skills(self, _mock_spacy_model, _mock_phrase_matcher_class):
        from .linkedin_profile_importer import build_linkedin_profile_import

        parsed = build_linkedin_profile_import(
            'Contact\nwww.linkedin.com/in/dev-profile\nTop Skills\nTechnical Standards\n'
            'Dev Candidate\nSoftware Engineer | Java • Kubernetes • AWS |\nMalaysia\nSummary\n'
            'Building backend services.\nExperience\nExample Co\nEngineer\n'
            'January 2024 - Present (6 months)'
        )

        self.assertEqual(parsed['skills'][0], 'Technical Standards')
        self.assertIn('Java', parsed['skills'])
        self.assertIn('Kubernetes', parsed['skills'])
        self.assertIn('AWS', parsed['skills'])

    @patch('apps.ai_services.skill_extractor._get_phrase_matcher_class', return_value=_FakePhraseMatcher)
    @patch('apps.ai_services.skill_extractor._load_spacy_model', return_value=_FakeNLP())
    def test_linkedin_profile_parser_handles_missing_optional_sections(self, _mock_spacy_model, _mock_phrase_matcher_class):
        from .linkedin_profile_importer import build_linkedin_profile_import

        parsed = build_linkedin_profile_import('Alex Applicant\nBackend Developer\nMalaysia\nExperience\nExample Co\nEngineer\nJanuary 2024 - Present (6 months)')

        self.assertEqual(parsed['full_name'], 'Alex Applicant')
        self.assertEqual(parsed['headline'], 'Backend Developer')
        self.assertEqual(parsed['location'], 'Malaysia')
        self.assertEqual(parsed['skills'], [])
        self.assertEqual(parsed['certifications'], [])
        self.assertEqual(parsed['education'], [])
        self.assertEqual(parsed['experience'][0]['company_name'], 'Example Co')



class ResumeMatchModelTests(SimpleTestCase):
    def test_ml_screening_requires_trained_artifact(self):
        from .ml.resume_matcher import build_ml_screening_result

        with self.assertRaises(AIServiceUnavailable):
            build_ml_screening_result(
                semantic_score=80,
                skill_score=75,
                experience_score=70,
                education_score=100,
                rule_based_score=78,
                matched_skills=['python', 'django'],
                missing_skills=['postgresql'],
                experience_gap={'gap_years': 0},
                education_gap={'gap_levels': 0},
                resume_text='Python Django developer with five years experience',
                job_text='Backend developer requiring Python Django PostgreSQL',
            )

    def test_score_to_label_uses_expected_match_bands(self):
        from .ml.resume_matcher import score_to_label

        self.assertEqual(score_to_label(90), 'strong_match')
        self.assertEqual(score_to_label(70), 'moderate_match')
        self.assertEqual(score_to_label(50), 'weak_match')
        self.assertEqual(score_to_label(30), 'not_suitable')

class ResumeContentValidationTests(SimpleTestCase):
    def test_valid_resume_with_skills_education_and_work_experience(self):
        from .resume_validation import validate_resume_text
        text = 'Skills: Python Django SQL. Education: Bachelor Degree in Computer Science from Example University. Experience: Worked as developer at Example Company for 2 years building APIs.'
        result = validate_resume_text(text)
        self.assertTrue(result['is_valid'])

    def test_valid_fresh_graduate_resume_with_projects(self):
        from .resume_validation import validate_resume_text
        text = 'Skills: Python React SQL. Education: Bachelor Degree in Information Technology, CGPA 3.7. Projects: Developed a mobile app and API dashboard system for final year project.'
        result = validate_resume_text(text)
        self.assertTrue(result['is_valid'])
        self.assertTrue(result['detected_fields']['projects'])

    def test_invalid_resume_missing_skills(self):
        from .resume_validation import validate_resume_text
        text = 'Education: Bachelor Degree in Business at University. Experience: Worked as executive at Example Company for 2 years managing operations and reports.'
        result = validate_resume_text(text)
        self.assertFalse(result['is_valid'])
        self.assertIn('skills', result['missing_fields'])

    def test_retail_sales_resume_detects_top_skills(self):
        from .resume_validation import validate_resume_text
        text = 'Top Skills: Attention To Detail, Accounting, Sales Target Management. Experience: Senior Sales Associate at company for 4 years. Education: Bachelor degree in Economics from University.'
        result = validate_resume_text(text)
        self.assertTrue(result['is_valid'])
        self.assertIn('accounting', result['detected_fields']['skills'])
        self.assertIn('sales target management', result['detected_fields']['skills'])

    def test_invalid_resume_missing_education(self):
        from .resume_validation import validate_resume_text
        text = 'Skills: Python Django SQL. Experience: Worked as developer at Example Company for 2 years building APIs and dashboard systems.'
        result = validate_resume_text(text)
        self.assertFalse(result['is_valid'])
        self.assertIn('education', result['missing_fields'])

    def test_invalid_resume_missing_experience_internship_or_projects(self):
        from .resume_validation import validate_resume_text
        text = 'Skills: Python Django SQL. Education: Bachelor Degree in Computer Science from Example University with CGPA 3.5. Professional summary and interests.'
        result = validate_resume_text(text)
        self.assertFalse(result['is_valid'])
        self.assertIn('experience_or_projects', result['missing_fields'])

    def test_unreadable_or_empty_resume_text(self):
        from .resume_validation import validate_resume_text
        result = validate_resume_text('')
        self.assertFalse(result['is_valid'])
        self.assertIn('Resume text could not be extracted or is too short for screening.', result['warnings'])

from apps.ai_services.speaker_diarization import (
    DIARIZATION_STATUS_NOT_CONFIGURED,
    DIARIZATION_STATUS_UNAVAILABLE,
    DiarizationUnavailable,
    _extract_speaker_turns,
    _format_diarization_error,
    _is_pyannote_sample_count_mismatch_error,
    _is_torchcodec_audio_read_error,
    _load_pyannote_pipeline,
    _run_diarization_pipeline,
    align_transcript_segments_to_speakers,
    apply_role_mapping,
    format_speaker_labelled_transcript,
    map_speakers_to_roles,
)
from apps.ai_services.transcription_service import TranscriptionUnavailable, build_speaker_aware_transcript_payload
from apps.ai_services.summary_service import (
    SummaryGenerationUnavailable,
    _parse_summary_content,
    run_real_summary,
)


class InterviewSpeakerDiarizationTests(SimpleTestCase):
    def test_extract_speaker_turns_supports_new_diarize_output_shape(self):
        turn = SimpleNamespace(start=0.5, end=2.25)
        diarization = SimpleNamespace(speaker_diarization=[(turn, 0)])

        self.assertEqual(
            _extract_speaker_turns(diarization),
            [{'speaker_id': 'SPEAKER_00', 'start_time': 0.5, 'end_time': 2.25}],
        )

    def test_extract_speaker_turns_supports_legacy_annotation_itertracks_shape(self):
        turn = SimpleNamespace(start=1.0, end=3.0)

        class FakeAnnotation:
            def itertracks(self, yield_label=False):
                self.yield_label = yield_label
                return [(turn, 'track', 'SPEAKER_01')]

        self.assertEqual(
            _extract_speaker_turns(FakeAnnotation()),
            [{'speaker_id': 'SPEAKER_01', 'start_time': 1.0, 'end_time': 3.0}],
        )

    def test_required_diarization_raises_instead_of_saving_plain_fallback(self):
        with patch.dict('os.environ', {'REQUIRE_SPEAKER_DIARIZATION': 'True'}):
            with self.assertRaisesMessage(TranscriptionUnavailable, 'Speaker diarization is required but did not complete'):
                build_speaker_aware_transcript_payload(
                    plain_transcript='Plain transcript text.',
                    transcript_segments=[],
                    audio_file=SimpleNamespace(name='interview.wav'),
                    metadata={'provider': 'local_whisper'},
                )

    def test_required_diarization_allows_completed_payload(self):
        with patch.dict('os.environ', {'REQUIRE_SPEAKER_DIARIZATION': 'True'}), patch(
            'apps.ai_services.transcription_service.run_speaker_diarization',
            return_value=[
                {'speaker_id': 'SPEAKER_00', 'start_time': 0.0, 'end_time': 1.0},
                {'speaker_id': 'SPEAKER_01', 'start_time': 1.0, 'end_time': 3.0},
            ],
        ):
            payload = build_speaker_aware_transcript_payload(
                plain_transcript='What did you build? I built APIs.',
                transcript_segments=[
                    {'start_time': 0.0, 'end_time': 1.0, 'text': 'What did you build?'},
                    {'start_time': 1.0, 'end_time': 3.0, 'text': 'I built APIs.'},
                ],
                audio_file=SimpleNamespace(name='interview.wav'),
                metadata={'provider': 'local_whisper'},
            )

        self.assertEqual(payload['transcript_json']['diarization_status'], 'completed')
        self.assertIn('Interviewer:', payload['transcript_text'])

    def test_detects_torchcodec_audio_read_error(self):
        error = RuntimeError('torchcodec is not available. Cannot read audio file.')

        self.assertTrue(_is_torchcodec_audio_read_error(error))
        self.assertFalse(_is_torchcodec_audio_read_error(RuntimeError('other failure')))


    def test_detects_pyannote_sample_count_mismatch_error(self):
        error = ValueError(
            'requested chunk [ 00:00:00.000 --> 00:00:10.000] from interview file resulted in '
            '439895 samples instead of the expected 441000 samples'
        )

        self.assertTrue(_is_pyannote_sample_count_mismatch_error(error))
        self.assertFalse(_is_pyannote_sample_count_mismatch_error(ValueError('other failure')))

    def test_run_diarization_pipeline_falls_back_to_waveform_input_when_sample_counts_mismatch(self):
        calls = []
        waveform_input = {'waveform': 'tensor', 'sample_rate': 16000}

        def fake_pipeline(audio_input):
            calls.append(audio_input)
            if isinstance(audio_input, str):
                raise ValueError(
                    'requested chunk [ 00:00:00.000 --> 00:00:10.000] from interview file resulted in '
                    '439895 samples instead of the expected 441000 samples'
                )
            return 'diarization'

        with patch('apps.ai_services.speaker_diarization._load_audio_waveform_with_whisper', return_value=waveform_input) as loader:
            result = _run_diarization_pipeline(fake_pipeline, 'interview.mp3')

        self.assertEqual(result, 'diarization')
        self.assertEqual(calls, ['interview.mp3', waveform_input])
        loader.assert_called_once_with('interview.mp3')

    def test_run_diarization_pipeline_falls_back_to_waveform_input_when_torchcodec_is_missing(self):
        calls = []
        waveform_input = {'waveform': 'tensor', 'sample_rate': 16000}

        def fake_pipeline(audio_input):
            calls.append(audio_input)
            if isinstance(audio_input, str):
                raise RuntimeError('torchcodec is not available. Cannot read audio file.')
            return 'diarization'

        with patch('apps.ai_services.speaker_diarization._load_audio_waveform_with_whisper', return_value=waveform_input) as loader:
            result = _run_diarization_pipeline(fake_pipeline, 'interview.mp3')

        self.assertEqual(result, 'diarization')
        self.assertEqual(calls, ['interview.mp3', waveform_input])
        loader.assert_called_once_with('interview.mp3')

    def test_diarization_unavailable_carries_fallback_status(self):
        not_configured = DiarizationUnavailable('disabled', status=DIARIZATION_STATUS_NOT_CONFIGURED)
        unavailable = DiarizationUnavailable('missing dependency', status=DIARIZATION_STATUS_UNAVAILABLE)

        self.assertEqual(not_configured.status, 'not_configured')
        self.assertEqual(unavailable.status, 'unavailable')

    def test_fallback_payload_uses_diarization_exception_status(self):
        with patch(
            'apps.ai_services.transcription_service.run_speaker_diarization',
            side_effect=DiarizationUnavailable('missing dependency', status=DIARIZATION_STATUS_UNAVAILABLE),
        ):
            payload = build_speaker_aware_transcript_payload(
                plain_transcript='Plain transcript text.',
                transcript_segments=[{'start_time': 0.0, 'end_time': 1.0, 'text': 'Plain transcript text.'}],
                audio_file=SimpleNamespace(name='interview.wav'),
                metadata={'provider': 'local_whisper'},
            )

        self.assertEqual(payload['transcript_text'], 'Plain transcript text.')
        self.assertEqual(payload['transcript_json']['diarization_status'], 'unavailable')
        self.assertEqual(payload['transcript_json']['diarization_warning'], 'missing dependency')

    def test_load_pyannote_pipeline_uses_token_keyword_for_current_versions(self):
        calls = []

        class FakePipeline:
            @staticmethod
            def from_pretrained(model_name, **kwargs):
                calls.append((model_name, kwargs))
                return 'pipeline'

        pyannote_audio = SimpleNamespace(Pipeline=FakePipeline)

        self.assertEqual(_load_pyannote_pipeline(pyannote_audio, 'pyannote/model', 'hf-token'), 'pipeline')
        self.assertEqual(calls, [('pyannote/model', {'token': 'hf-token'})])

    def test_load_pyannote_pipeline_falls_back_to_legacy_use_auth_token_keyword(self):
        calls = []

        class FakePipeline:
            @staticmethod
            def from_pretrained(model_name, **kwargs):
                calls.append((model_name, kwargs))
                if 'token' in kwargs:
                    raise TypeError("from_pretrained() got an unexpected keyword argument 'token'")
                return 'pipeline'

        pyannote_audio = SimpleNamespace(Pipeline=FakePipeline)

        self.assertEqual(_load_pyannote_pipeline(pyannote_audio, 'pyannote/model', 'hf-token'), 'pipeline')
        self.assertEqual(calls[-1], ('pyannote/model', {'use_auth_token': 'hf-token'}))

    def test_diarization_error_includes_type_and_message(self):
        error = _format_diarization_error(TypeError("from_pretrained() got an unexpected keyword argument 'use_auth_token'"))

        self.assertIn('TypeError', error)
        self.assertIn('use_auth_token', error)

    def test_formats_structured_speaker_segments_and_merges_consecutive_roles(self):
        formatted = format_speaker_labelled_transcript([
            {'speaker_id': 'SPEAKER_00', 'role': 'Interviewer', 'start_time': 0.0, 'end_time': 1.0, 'text': 'Good afternoon.'},
            {'speaker_id': 'SPEAKER_00', 'role': 'Interviewer', 'start_time': 1.0, 'end_time': 2.0, 'text': 'Please have a seat.'},
            {'speaker_id': 'SPEAKER_01', 'role': 'Candidate', 'start_time': 2.0, 'end_time': 3.0, 'text': 'Thank you.'},
        ])

        self.assertEqual(formatted, 'Interviewer: Good afternoon. Please have a seat.\n\nCandidate: Thank you.')
        self.assertNotIn('Interviewee', formatted)

    def test_maps_question_asking_speaker_to_interviewer_and_other_to_candidate(self):
        aligned_segments = [
            {'speaker_id': 'SPEAKER_00', 'start_time': 0.0, 'end_time': 2.0, 'text': 'Can you explain your Django experience?'},
            {'speaker_id': 'SPEAKER_00', 'start_time': 6.0, 'end_time': 8.0, 'text': 'What projects did you lead?'},
            {'speaker_id': 'SPEAKER_01', 'start_time': 2.0, 'end_time': 6.0, 'text': 'I built several APIs and led a final year project.'},
        ]

        mapping, warning = map_speakers_to_roles(aligned_segments)
        self.assertIsNone(warning)
        self.assertEqual(mapping['SPEAKER_00'], 'Interviewer')
        self.assertEqual(mapping['SPEAKER_01'], 'Candidate')
        self.assertNotIn('Interviewee', mapping.values())

    def test_aligns_transcript_segment_to_largest_timestamp_overlap(self):
        aligned = align_transcript_segments_to_speakers(
            [{'start_time': 1.0, 'end_time': 4.0, 'text': 'I worked with Django.'}],
            [
                {'speaker_id': 'SPEAKER_00', 'start_time': 0.0, 'end_time': 1.5},
                {'speaker_id': 'SPEAKER_01', 'start_time': 1.5, 'end_time': 4.0},
            ],
        )

        self.assertEqual(aligned[0]['speaker_id'], 'SPEAKER_01')

    def test_fallback_payload_keeps_plain_transcript_when_diarization_unavailable(self):
        with patch('apps.ai_services.transcription_service.run_speaker_diarization', side_effect=Exception('missing diarization')):
            payload = build_speaker_aware_transcript_payload(
                plain_transcript='Plain transcript text.',
                transcript_segments=[{'start_time': 0.0, 'end_time': 1.0, 'text': 'Plain transcript text.'}],
                audio_file=SimpleNamespace(name='interview.wav'),
                metadata={'provider': 'local_whisper'},
            )

        self.assertEqual(payload['transcript_text'], 'Plain transcript text.')
        self.assertEqual(payload['transcript_json']['diarization_status'], 'failed')
        self.assertTrue(payload['transcript_json']['diarization_warning'])
        self.assertEqual(payload['transcript_json']['segments'], [])

    def test_completed_payload_saves_readable_and_structured_transcript(self):
        with patch(
            'apps.ai_services.transcription_service.run_speaker_diarization',
            return_value=[
                {'speaker_id': 'SPEAKER_00', 'start_time': 0.0, 'end_time': 2.0},
                {'speaker_id': 'SPEAKER_01', 'start_time': 2.0, 'end_time': 5.0},
            ],
        ):
            payload = build_speaker_aware_transcript_payload(
                plain_transcript='Can you explain your work? I built APIs.',
                transcript_segments=[
                    {'start_time': 0.0, 'end_time': 2.0, 'text': 'Can you explain your work?'},
                    {'start_time': 2.0, 'end_time': 5.0, 'text': 'I built APIs.'},
                ],
                audio_file=SimpleNamespace(name='interview.wav'),
                metadata={'provider': 'local_whisper'},
            )

        self.assertIn('Interviewer:', payload['transcript_text'])
        self.assertIn('Candidate:', payload['transcript_text'])
        self.assertEqual(payload['transcript_json']['diarization_status'], 'completed')
        self.assertEqual(len(payload['transcript_json']['segments']), 2)


class InterviewSummaryGeminiTests(SimpleTestCase):
    def _summary_json(self):
        return json.dumps({
            'strengths': 'Clear examples of relevant project experience.',
            'weaknesses': 'Could provide more measurable impact details.',
            'communication_score': 8,
            'overall_impression': 'Professional and concise interview responses.',
            'editable_summary_text': 'Candidate communicated relevant experience clearly.',
        })

    def test_summary_parser_accepts_json_embedded_in_provider_text(self):
        content = """Here is the structured summary:
```json
{
  "strengths": "Clear API examples.",
  "weaknesses": "Needs more metrics.",
  "communication_score": 8,
  "overall_impression": "Professional responses.",
  "editable_summary_text": "Candidate gave clear API examples."
}
```
Please review before saving."""

        parsed = _parse_summary_content(content)

        self.assertEqual(parsed['strengths'], 'Clear API examples.')
        self.assertEqual(parsed['communication_score'], 8)

    def test_real_summary_can_use_gemini_provider(self):
        with patch.dict('os.environ', {
            'USE_REAL_SUMMARY': 'True',
            'SUMMARY_PROVIDER': 'gemini',
            'GEMINI_API_KEY': 'test-gemini-key',
            'SUMMARY_MODEL': 'gemini-3.5-flash',
        }), patch(
            'apps.ai_services.summary_service._call_gemini_summary',
            return_value=self._summary_json(),
        ) as gemini_call:
            payload = run_real_summary('Interviewer: Tell me about your projects. Candidate: I built APIs.')

        gemini_call.assert_called_once()
        self.assertEqual(payload['summary_json']['provider'], 'gemini')
        self.assertEqual(payload['summary_json']['model'], 'gemini-3.5-flash')
        self.assertEqual(payload['communication_score'], 8)
        self.assertNotIn('final hiring decision', payload['editable_summary_text'].lower())

    def test_retired_gemini_model_is_replaced_with_current_default(self):
        with patch.dict('os.environ', {
            'USE_REAL_SUMMARY': 'True',
            'SUMMARY_PROVIDER': 'gemini',
            'GEMINI_API_KEY': 'test-gemini-key',
            'SUMMARY_MODEL': 'gemini-2.5-flash',
        }), patch(
            'apps.ai_services.summary_service._call_gemini_summary',
            return_value=self._summary_json(),
        ) as gemini_call:
            payload = run_real_summary('Transcript text')

        self.assertEqual(gemini_call.call_args.args[2], 'gemini-3.5-flash')
        self.assertEqual(payload['summary_json']['model'], 'gemini-3.5-flash')

    def test_gemini_summary_requires_gemini_api_key(self):
        with patch.dict('os.environ', {
            'USE_REAL_SUMMARY': 'True',
            'SUMMARY_PROVIDER': 'gemini',
            'GEMINI_API_KEY': '',
        }, clear=False):
            with self.assertRaisesMessage(SummaryGenerationUnavailable, 'GEMINI_API_KEY is required'):
                run_real_summary('Transcript text')

    def test_gemini_provider_failure_includes_client_error_details(self):
        class ClientError(Exception):
            message = 'API key not valid. Please pass a valid API key.'
            code = 400

        with patch.dict('os.environ', {
            'USE_REAL_SUMMARY': 'True',
            'SUMMARY_PROVIDER': 'gemini',
            'GEMINI_API_KEY': 'bad-key',
            'SUMMARY_MODEL': 'gemini-3.5-flash',
        }), patch(
            'apps.ai_services.summary_service._call_gemini_summary',
            side_effect=ClientError('API key not valid. Please pass a valid API key.'),
        ):
            with self.assertRaises(SummaryGenerationUnavailable) as error_context:
                run_real_summary('Transcript text')

        error_message = str(error_context.exception)
        self.assertIn('Gemini summary generation failed for model gemini-3.5-flash', error_message)
        self.assertIn('ClientError', error_message)
        self.assertIn('API key not valid', error_message)
