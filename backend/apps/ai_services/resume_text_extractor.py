"""Local resume text extraction helpers for supported document formats."""

from importlib.util import find_spec
from pathlib import Path


SUPPORTED_RESUME_EXTENSIONS = {'.pdf', '.docx'}


class ResumeTextExtractionError(ValueError):
    """Raised when a local resume file cannot be converted to text."""


def extract_text_from_pdf(file_path):
    """Extract text from each page of a local PDF file."""
    path = _validate_local_file(file_path, expected_extension='.pdf')
    _require_dependency('fitz', 'PyMuPDF', 'PDF')
    import fitz

    try:
        with fitz.open(path) as document:
            return _clean_text('\n'.join(page.get_text() for page in document))
    except Exception as exc:
        raise ResumeTextExtractionError(f'Unable to extract text from PDF file: {path.name}') from exc


def extract_text_from_docx(file_path):
    """Extract text from paragraphs and table cells in a local DOCX file."""
    path = _validate_local_file(file_path, expected_extension='.docx')
    _require_dependency('docx', 'python-docx', 'DOCX')
    from docx import Document

    try:
        document = Document(path)
        text_parts = [paragraph.text for paragraph in document.paragraphs]
        text_parts.extend(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        return _clean_text('\n'.join(text_parts))
    except Exception as exc:
        raise ResumeTextExtractionError(f'Unable to extract text from DOCX file: {path.name}') from exc


def extract_resume_text(file_path):
    """Extract text from a supported local resume file based on its extension."""
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == '.pdf':
        return extract_text_from_pdf(path)
    if extension == '.docx':
        return extract_text_from_docx(path)

    supported_extensions = ', '.join(sorted(SUPPORTED_RESUME_EXTENSIONS))
    raise ResumeTextExtractionError(
        f'Unsupported resume file type: {extension or "no extension"}. '
        f'Supported file types are: {supported_extensions}.'
    )


def _require_dependency(module_name, package_name, document_type):
    """Keep optional document libraries from blocking Django application startup."""
    if find_spec(module_name) is None:
        raise ResumeTextExtractionError(
            f'{document_type} resume extraction requires {package_name}. '
            f'Install backend dependencies with `pip install -r requirements.txt`.'
        )


def _validate_local_file(file_path, expected_extension):
    path = Path(file_path)

    if path.suffix.lower() != expected_extension:
        raise ResumeTextExtractionError(f'Expected a {expected_extension} file: {path.name}')
    if not path.is_file():
        raise ResumeTextExtractionError(f'Resume file does not exist: {path}')

    return path


def _clean_text(text):
    return '\n'.join(line.strip() for line in text.splitlines() if line.strip())
